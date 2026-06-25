# -*- coding: utf-8 -*-
"""Render Block Model (xem knowledge/format-specs/00-...) thành tài liệu Word.

Dùng chung cho bài đọc, kịch bản video, bài tập. LLM chỉ sinh `blocks`, module này dựng .docx.
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CODE_FONT = "Consolas"
CODE_SHADE = "F2F2F2"  # nền xám nhạt cho code


def new_doc():
    doc = Document()
    # Font mặc định dễ đọc tiếng Việt
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    return doc


def _set_para_shading(paragraph, hex_fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    p_pr.append(shd)


def _add_code(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Pt(6)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)
    _set_para_shading(para, CODE_SHADE)
    lines = text.split("\n")
    for i, line in enumerate(lines):
        run = para.add_run(line if line else "")
        run.font.name = CODE_FONT
        run.font.size = Pt(10)
        # đảm bảo font áp cho cả East Asian
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:ascii"), CODE_FONT)
        rfonts.set(qn("w:hAnsi"), CODE_FONT)
        if i < len(lines) - 1:
            run.add_break()
    return para


def _add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for j, h in enumerate(headers):
        hdr[j].text = str(h)
        for p in hdr[j].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for j in range(len(headers)):
            cells[j].text = str(row[j]) if j < len(row) else ""
    return table


def add_blocks(doc, blocks):
    """Duyệt và render danh sách block vào doc."""
    for b in blocks:
        t = b.get("type", "p")
        if t == "h1":
            doc.add_heading(b.get("text", ""), level=1)
        elif t == "h2":
            doc.add_heading(b.get("text", ""), level=2)
        elif t == "h3":
            doc.add_heading(b.get("text", ""), level=3)
        elif t in ("p", "narration"):
            doc.add_paragraph(b.get("text", ""))
        elif t == "quote":
            para = doc.add_paragraph()
            run = para.add_run(b.get("text", ""))
            run.italic = True
        elif t == "marker":
            para = doc.add_paragraph()
            txt = b.get("text", "").strip()
            if not (txt.startswith("[") and txt.endswith("]")):
                txt = "[" + txt.strip("[]") + "]"
            run = para.add_run(txt)
            run.bold = True
        elif t == "bullets":
            for it in b.get("items", []):
                doc.add_paragraph(str(it), style="List Bullet")
        elif t == "numbers":
            for it in b.get("items", []):
                doc.add_paragraph(str(it), style="List Number")
        elif t == "code":
            _add_code(doc, b.get("text", ""))
        elif t == "table":
            _add_table(doc, b.get("headers", []), b.get("rows", []))
        else:
            doc.add_paragraph(b.get("text", ""))
    return doc
