# -*- coding: utf-8 -*-
"""Kiểm tra hợp đồng định dạng trên bộ artifact đã render của một session.

Chạy:  PYTHONIOENCODING=utf-8 python tools/qa_check.py "output/Session 02 - ..."
Thoát mã != 0 nếu có lỗi BLOCKER.
"""
import glob
import json
import os
import re
import sys
import zipfile

import openpyxl
from docx import Document
from docx.oxml.ns import qn

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
PLACEHOLDERS = ["todo", "lorem ipsum", "[điền", "placeholder", "xxx", "tbd"]
QUIZ_HEADERS = [
    "question_content", "answer_1", "explanation_answer_1", "answer_2", "explanation_answer_2",
    "answer_3", "explanation_answer_3", "answer_4", "explanation_answer_4",
    "isCorrect", "difficulty", "category",
]

errors = []   # BLOCKER
warns = []    # NÊN SỬA


def err(msg):
    errors.append(msg)


def warn(msg):
    warns.append(msg)


def doc_paras(path):
    return [p for p in Document(path).paragraphs]


def doc_text(path):
    return "\n".join(p.text for p in Document(path).paragraphs)


def headings(paras, level_name):
    return [p.text for p in paras if p.style and p.style.name == level_name]


def has_code(path):
    for p in Document(path).paragraphs:
        for r in p.runs:
            rpr = r._element.find(qn("w:rPr"))
            if rpr is not None:
                rfonts = rpr.find(qn("w:rFonts"))
                if rfonts is not None and "Consolas" in (rfonts.get(qn("w:ascii")) or ""):
                    return True
    return False


def check_placeholders(path, text, phs=None):
    low = text.lower()
    for ph in (phs if phs is not None else PLACEHOLDERS):
        if ph in low:
            warn("%s: còn placeholder nghi vấn '%s'" % (os.path.basename(path), ph))


# ---------- per-artifact checks ----------

def check_readings(d):
    files = glob.glob(os.path.join(d, "Bài đọc", "*.docx"))
    if not files:
        err("Thiếu thư mục/đầu ra Bài đọc")
        return
    for f in files:
        if os.path.basename(f).startswith("~$"):
            continue
        paras = doc_paras(f)
        h1 = headings(paras, "Heading 1")
        h2 = headings(paras, "Heading 2")
        name = os.path.basename(f)
        if len(h1) < 1:
            err("Bài đọc %s: thiếu Heading 1" % name)
        numbered = [h for h in h2 if re.match(r"^\s*\d+\.", h)]
        if len(numbered) < 3:
            warn("Bài đọc %s: <3 mục h2 đánh số (%d)" % (name, len(numbered)))
        if not any("tổng kết" in h.lower() for h in h2):
            warn("Bài đọc %s: thiếu mục 'Tổng Kết'" % name)
        if not any("tham khảo" in h.lower() for h in h2):
            warn("Bài đọc %s: thiếu mục 'Tài Liệu Tham Khảo'" % name)
        if not has_code(f):
            warn("Bài đọc %s: không phát hiện khối code" % name)
        check_placeholders(f, doc_text(f))


def check_videos(d, expected_lessons):
    files = [f for f in glob.glob(os.path.join(d, "Kịch bản quay video", "*.docx"))
             if not os.path.basename(f).startswith("~$")]
    if not files:
        err("Thiếu thư mục/đầu ra Kịch bản quay video")
        return
    if expected_lessons and len(files) < expected_lessons:
        warn("Kịch bản video: có %d file, khung CT có %d Lesson" % (len(files), expected_lessons))
    for f in files:
        txt = doc_text(f)
        name = os.path.basename(f)
        if "rikkei education" not in txt.lower():
            warn("Video %s: thiếu câu chào 'Rikkei Education'" % name)
        if txt.count("## ") < 2:
            warn("Video %s: <2 heading '## '" % name)
        if "tổng kết bài giảng" not in txt.lower():
            warn("Video %s: thiếu 'Tổng kết bài giảng'" % name)
        if "[chuyển tiếp slide]" not in txt.lower():
            warn("Video %s: thiếu marker [Chuyển tiếp slide]" % name)
        check_placeholders(f, txt)


def check_exercises(d):
    files = [f for f in glob.glob(os.path.join(d, "Bài tập", "*.docx"))
             if not os.path.basename(f).startswith("~$")]
    if not files:
        err("Thiếu thư mục/đầu ra Bài tập")
        return
    if len(files) < 6:
        warn("Bài tập: chỉ có %d bài (khuyến nghị ≥6)" % len(files))
    names = " ".join(os.path.basename(f).lower() for f in files)
    for need in ["vận dụng cơ bản", "phân tích", "sáng tạo"]:
        if need not in names:
            warn("Bài tập: thiếu cấp độ '%s'" % need)
    for f in files:
        txt = doc_text(f)
        name = os.path.basename(f)
        if "bối cảnh" not in txt.lower():
            warn("Bài tập %s: thiếu 'Bối cảnh nghiệp vụ'" % name)
        if "nộp bài" not in txt.lower():
            warn("Bài tập %s: thiếu 'Quy định nộp bài'" % name)
        # 'todo' là code-stub hợp lệ trong đề bài (sinh viên tự hoàn thiện) → bỏ qua riêng cho bài tập
        check_placeholders(f, txt, [p for p in PLACEHOLDERS if p != "todo"])


def check_quizzes(d):
    files = [f for f in glob.glob(os.path.join(d, "Câu hỏi Quiziz session", "*.xlsx"))
             if not os.path.basename(f).startswith("~$")]
    if not files:
        err("Thiếu thư mục/đầu ra Quiz")
        return
    has_dau = any("daugio" in os.path.basename(f).lower() for f in files)
    has_cuoi = any("cuoigio" in os.path.basename(f).lower() for f in files)
    if not has_dau:
        err("Thiếu Quiz Đầu giờ")
    if not has_cuoi:
        err("Thiếu Quiz Cuối giờ")
    for f in files:
        name = os.path.basename(f)
        wb = openpyxl.load_workbook(f, data_only=True)
        ws = wb.active
        header = [ws.cell(1, c).value for c in range(1, 13)]
        if header != QUIZ_HEADERS:
            err("Quiz %s: header 12 cột không khớp lược đồ" % name)
            continue
        rows = list(ws.iter_rows(min_row=2, values_only=True))
        cats = set()
        for i, r in enumerate(rows, 2):
            if r[0] is None or str(r[0]).strip() == "":
                continue
            ans = [r[1], r[3], r[5], r[7]]
            exps = [r[2], r[4], r[6], r[8]]
            if any(a is None or str(a).strip() == "" for a in ans):
                err("Quiz %s hàng %d: thiếu đáp án" % (name, i))
            if any(e is None or str(e).strip() == "" for e in exps):
                warn("Quiz %s hàng %d: thiếu giải thích" % (name, i))
            try:
                ic = int(r[9])
                if ic not in (1, 2, 3, 4):
                    err("Quiz %s hàng %d: isCorrect=%r ngoài {1..4}" % (name, i, r[9]))
            except (TypeError, ValueError):
                err("Quiz %s hàng %d: isCorrect không phải số" % (name, i))
            if r[11]:
                cats.add(str(r[11]).strip().upper())
        if "daugio" in name.lower():
            if "BÀI CŨ" not in cats or "BÀI MỚI" not in cats:
                warn("Quiz Đầu giờ %s: nên có cả BÀI CŨ và BÀI MỚI (có: %s)" % (name, sorted(cats)))


def check_slides(d):
    mds = glob.glob(os.path.join(d, "Bài giảng", "*.md"))
    if not mds:
        err("Thiếu outline bài giảng (.md)")
        return
    txt = open(mds[0], encoding="utf-8").read()
    low = txt.lower()
    for need, label in [("nội dung", "agenda NỘI DUNG"), ("tổng kết", "TỔNG KẾT"),
                        ("kết thúc", "KẾT THÚC")]:
        if need not in low:
            warn("Outline slide: thiếu slide %s" % label)
    if "speaker notes" not in low:
        warn("Outline slide: thiếu speaker notes")
    check_placeholders(mds[0], txt)


def check_mindmap(d):
    files = glob.glob(os.path.join(d, "Mindmap", "*.xmind"))
    if not files:
        err("Thiếu Mindmap (.xmind)")
        return
    try:
        z = zipfile.ZipFile(files[0])
        content = json.loads(z.read("content.json").decode("utf-8"))
        root = content[0]["rootTopic"]
        kids = root.get("children", {}).get("attached", [])
        if not kids:
            warn("Mindmap: root không có nhánh con")
    except Exception as e:
        err("Mindmap không hợp lệ: %s" % e)


def lessons_for_session(session_dir):
    fw = os.path.join(ROOT, "knowledge", "framework.json")
    if not os.path.exists(fw):
        return None
    m = re.search(r"Session\s*0*(\d+)", os.path.basename(session_dir.rstrip("/\\")))
    if not m:
        return None
    no = int(m.group(1))
    data = json.load(open(fw, encoding="utf-8"))
    for s in data["sessions"]:
        if s.get("session_no") == no:
            return len([l for l in s["lessons"] if l.get("no")])
    return None


def main():
    if len(sys.argv) < 2:
        print('Cách dùng: python tools/qa_check.py "<session_dir>"')
        sys.exit(2)
    d = sys.argv[1]
    if not os.path.isdir(d):
        print("Không thấy thư mục:", d)
        sys.exit(2)
    expected = lessons_for_session(d)
    check_readings(d)
    check_videos(d, expected)
    check_exercises(d)
    check_quizzes(d)
    check_slides(d)
    check_mindmap(d)

    print("=== KẾT QUẢ QA: %s ===" % os.path.basename(d.rstrip("/\\")))
    if errors:
        print("\nBLOCKER (%d):" % len(errors))
        for e in errors:
            print("  [X]", e)
    if warns:
        print("\nNÊN SỬA (%d):" % len(warns))
        for w in warns:
            print("  [!]", w)
    if not errors and not warns:
        print("PASS toàn bộ — không phát hiện vấn đề.")
    elif not errors:
        print("\n=> PASS (không có BLOCKER; %d cảnh báo)" % len(warns))
    else:
        print("\n=> FAIL (%d BLOCKER)" % len(errors))
    sys.exit(1 if errors else 0)


if __name__ == "__main__":
    main()
