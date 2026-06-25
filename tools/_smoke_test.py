# -*- coding: utf-8 -*-
"""Smoke test: dựng spec tối thiểu cho cả 6 renderer, render, rồi mở lại kiểm tra."""
import os
import sys
import zipfile

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from lib.io_utils import dump_json
import render_all

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SDIR = os.path.join(ROOT, "output", "_SMOKE")
SPEC = os.path.join(SDIR, "_spec")

specs = {
    "reading_01.json": {
        "filename": "BÀI ĐỌC_ SMOKE", "subdir": "Bài đọc",
        "blocks": [
            {"type": "h1", "text": "Bài Đọc Chuyên Sâu: Kiểm thử"},
            {"type": "h2", "text": "1. Mục tiêu"},
            {"type": "p", "text": "Đoạn văn tiếng Việt có dấu."},
            {"type": "code", "lang": "python", "text": "from fastapi import FastAPI\napp = FastAPI()"},
            {"type": "bullets", "items": ["Ý một", "Ý hai"]},
            {"type": "h2", "text": "Tổng Kết"},
            {"type": "h2", "text": "Tài Liệu Tham Khảo"},
        ],
    },
    "video_01.json": {
        "filename": "Lesson 01 - Smoke", "subdir": "Kịch bản quay video", "lesson_no": 1,
        "blocks": [
            {"type": "h2", "text": "## Mở đầu"},
            {"type": "narration", "text": "Chào mừng các em đã quay trở lại với hệ thống Elearning của Rikkei Education."},
            {"type": "marker", "text": "[Chuyển tiếp slide]"},
            {"type": "h2", "text": "## Tổng kết bài giảng"},
            {"type": "narration", "text": "Cảm ơn các em đã theo dõi bài học hôm nay."},
        ],
    },
    "exercise_01.json": {
        "filename": "[Vận dụng cơ bản 1] Smoke", "subdir": "Bài tập", "level": "Vận dụng cơ bản",
        "blocks": [
            {"type": "h1", "text": "Bài tập smoke"},
            {"type": "h2", "text": "1. Bối cảnh nghiệp vụ"},
            {"type": "p", "text": "Bối cảnh."},
            {"type": "table", "headers": ["Tiêu chí", "PA1", "PA2"], "rows": [["Rõ ràng", "", ""]]},
        ],
        "submission": {"lop": "[Tên Lớp]", "mon": "FastAPI", "session": "Session02", "ex": "Ex01",
                       "example": "HNKS25CNTT1_FastAPI_Session02_Ex01"},
    },
    "quiz_daugio.json": {
        "filename": "Import_Quiz_DauGio_SMOKE", "subdir": "Câu hỏi Quiziz session",
        "sheet_name": "Quiz_DauGio_OOP",
        "questions": [
            {"q": "Từ khóa khai báo hàm Python?", "answers": ["define", "func", "function", "def"],
             "explanations": ["sai", "sai", "sai", "đúng"], "correct": 4, "difficulty": 4, "category": "BÀI CŨ"},
            {"q": "Uvicorn là gì?", "answers": ["WSGI", "ASGI server", "DB", "IDE"],
             "explanations": ["sai", "đúng", "sai", "sai"], "correct": 2, "difficulty": 8, "category": "BÀI MỚI"},
        ],
    },
    "slide_outline.json": {
        "filename": "Session SMOKE", "subdir": "Bài giảng", "title": "Smoke", "version": "1.0",
        "module": "Module: Test",
        "slides": [
            {"layout": "title", "title": "Smoke", "content": ["Session 02"]},
            {"layout": "bullets", "section": "1. A", "title": "Nội dung", "content": ["x", "y"],
             "diagram_hint": "sơ đồ", "speaker_notes": "khớp video lesson 01"},
            {"layout": "closing", "title": "KẾT THÚC", "content": ["HỌC VIỆN ĐÀO TẠO LẬP TRÌNH CHẤT LƯỢNG NHẬT BẢN"]},
        ],
    },
    "mindmap.json": {
        "filename": "Session SMOKE", "subdir": "Mindmap", "root": "Smoke session",
        "children": [{"title": "1. Lesson", "children": [{"title": "Khái niệm A"}]}],
    },
}

for name, obj in specs.items():
    dump_json(obj, os.path.join(SPEC, name))

render_all.render_all(SDIR)

# Verify
from docx import Document
import openpyxl
import json as _json

checks = []
checks.append(("reading docx", len(Document(os.path.join(SDIR, "Bài đọc", "BÀI ĐỌC_ SMOKE.docx")).paragraphs) > 0))
checks.append(("video docx", len(Document(os.path.join(SDIR, "Kịch bản quay video", "Lesson 01 - Smoke.docx")).paragraphs) > 0))
ex_doc = Document(os.path.join(SDIR, "Bài tập", "[Vận dụng cơ bản 1] Smoke.docx"))
checks.append(("exercise has submission", any("nộp bài" in p.text.lower() for p in ex_doc.paragraphs)))
wb = openpyxl.load_workbook(os.path.join(SDIR, "Câu hỏi Quiziz session", "Import_Quiz_DauGio_SMOKE.xlsx"))
ws = wb.active
checks.append(("quiz 12 cols", ws.max_column == 12))
checks.append(("quiz header", ws.cell(1, 1).value == "question_content" and ws.cell(1, 10).value == "isCorrect"))
checks.append(("slide md exists", os.path.exists(os.path.join(SDIR, "Bài giảng", "Session SMOKE.md"))))
xm = os.path.join(SDIR, "Mindmap", "Session SMOKE.xmind")
z = zipfile.ZipFile(xm)
content = _json.loads(z.read("content.json").decode("utf-8"))
checks.append(("xmind root", content[0]["rootTopic"]["title"] == "Smoke session"))
checks.append(("xmind child", content[0]["rootTopic"]["children"]["attached"][0]["title"] == "1. Lesson"))

print("\n--- CHECKS ---")
ok = True
for name, res in checks:
    print(("PASS" if res else "FAIL"), name)
    ok = ok and res
print("ALL PASS" if ok else "SOME FAILED")
sys.exit(0 if ok else 1)
